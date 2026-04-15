"""Runner-level report generation: Markdown and DOCX outputs."""

import html
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Sequence

from jamovi_runner.formatting import render_markdown_table_block


def normalize_output_stem(raw_value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", raw_value.strip())
    cleaned = cleaned.strip("-.")
    return cleaned or "jamovi-project"


def resolve_output_paths(
    data_path: Path,
    output_dir: str | None,
    output_basename: str | None,
    spec_output_basename: str | None,
) -> tuple[Path, Path, Path, Path]:
    destination = Path(output_dir) if output_dir else data_path.parent
    destination.mkdir(parents=True, exist_ok=True)
    stem_source = output_basename or spec_output_basename or f"{data_path.stem}-jamovi-project"
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    stem = normalize_output_stem(stem_source)
    base = destination / f"{stem}-{timestamp}"
    return (
        base.with_suffix(".omv"),
        base.with_suffix(".md"),
        base.with_suffix(".docx"),
        base,
    )


def format_variables_for_markdown(variables: dict[str, Any]) -> str:
    parts: list[str] = []
    for key, value in variables.items():
        if isinstance(value, list):
            parts.append(f"{key}={', '.join(str(v) for v in value)}")
        elif isinstance(value, dict):
            inner = ", ".join(f"{k}={v}" for k, v in value.items())
            parts.append(f"{key}={{{inner}}}")
        else:
            parts.append(f"{key}={value}")
    return "; ".join(parts)


def build_runner_markdown_report(
    data_path: Path,
    omv_path: Path | None,
    parse_contract: dict[str, Any] | None,
    records: Sequence[Any],
    teardown_errors: Sequence[str],
    mode: str,
    column_manifest: dict[str, Any] | None = None,
    timings: dict[str, float] | None = None,
    table_style: str = "gfm",
    sidecar_info: dict[str, Any] | None = None,
) -> str:
    lines = [
        "# Jamovi Project Summary",
        "",
        f"- Generated: {datetime.now().isoformat(timespec='seconds')}",
        f"- Mode: {mode}",
        f"- Data file: `{data_path}`",
        f"- OMV path: `{omv_path}`" if omv_path else "- OMV path: not saved",
        f"- Table style: `{table_style}`",
        "",
    ]
    table_index = 1

    if timings:
        lines.append(f"- Total Time: {timings.get('total_seconds', 0.0):.2f}s")
        phase_labels = [
            ("preprocess_seconds", "Preprocess"),
            ("open_seconds", "Open"),
            ("run_seconds", "Run"),
            ("save_seconds", "Save"),
        ]
        phase_parts = []
        for key, label in phase_labels:
            val = timings.get(key)
            if val is not None:
                phase_parts.append(f"{label} {val:.2f}s")
        if phase_parts:
            lines.append(f"  - Phases: {' | '.join(phase_parts)}")
        lines.append("")

    if column_manifest:
        lines.append("## Column Mapping")
        lines.append("")
        lines.append("Data was preprocessed. The following aliases were used internally:")
        lines.append("")
        table_rows_block = [{"Clean Alias": f"`{alias}`", "Original Header": original} for alias, original in column_manifest.items()]
        table_lines, table_index = render_markdown_table_block(
            table_rows_block,
            "Column Mapping",
            table_style,
            table_index,
        )
        lines.extend(table_lines)
        lines.append("")

    if sidecar_info and sidecar_info.get("dir"):
        lines.append("## Preprocess Artifacts")
        lines.append("")
        lines.append(f"- Output directory: `{sidecar_info['dir']}`")
        if sidecar_info.get("analysis_ready"):
            lines.append(f"- Analysis-ready CSV: `{sidecar_info['analysis_ready']}`")
        if sidecar_info.get("column_manifest"):
            lines.append(f"- Column manifest: `{sidecar_info['column_manifest']}`")
        lines.append("")

    if parse_contract is not None:
        lines.extend(
            [
                "## Parse Contract",
                "",
                f"- is_executable: `{parse_contract['is_executable']}`",
                f"- missing_info: {parse_contract['missing_info'] or 'n/a'}",
                "",
            ]
        )

    successful = [record for record in records if record.status == "success"]
    failed = [record for record in records if record.status != "success"]

    lines.extend(["## Successful Analyses", ""])
    if not successful:
        lines.append("No analyses completed successfully.")
        lines.append("")
    else:
        for index, record in enumerate(successful, start=1):
            lines.append(f"### {index}. {record.title} (`{record.analysis_type}`)")
            lines.append("")
            lines.append(f"- Variables: {format_variables_for_markdown(record.variables)}")
            lines.append(f"- Status: {record.status_detail}")
            if record.note:
                lines.append(f"- Note: {record.note}")
            lines.append("")
            if record.summary_sections:
                if table_style == "apa":
                    from jamovi_runner.report import build_markdown_report as _build_markdown_report_from_package

                    apa_lines, table_index = _build_markdown_report_from_package(
                        record.summary_sections, table_style, table_index
                    )
                    lines.extend(apa_lines)
                    lines.append("")
                else:
                    for section in record.summary_sections:
                        lines.append(f"#### {section['title']}")
                        table_lines, table_index = render_markdown_table_block(
                            section["rows"],
                            section["title"],
                            table_style,
                            table_index,
                        )
                        lines.extend(table_lines)
                        lines.append("")
            else:
                lines.append("Key results could not be extracted automatically. Open the `.omv` file for the full result tree.")
                lines.append("")

    lines.extend(["## Failed Analyses", ""])
    if not failed:
        lines.append("No analysis failures were recorded.")
        lines.append("")
    else:
        failure_rows = [
            {"Analysis": record.analysis_type, "Failure Type": record.status, "Reason": record.status_detail}
            for record in failed
        ]
        table_lines, table_index = render_markdown_table_block(
            failure_rows,
            "Failed Analyses",
            table_style,
            table_index,
        )
        lines.extend(table_lines)
        lines.append("")

    lines.extend(["## Teardown", ""])
    if teardown_errors:
        for issue in teardown_errors:
            lines.append(f"- {issue}")
    else:
        lines.append("- No teardown issues recorded.")
    lines.append("")
    return "\n".join(lines)


def render_markdown_html(markdown_text: str) -> tuple[str, str | None]:
    try:
        import markdown  # type: ignore
    except Exception as exc:
        escaped = html.escape(markdown_text)
        fallback = f"<html><head><meta charset=\"utf-8\"></head><body><pre>{escaped}</pre></body></html>"
        return fallback, f"HTML export fallback used (markdown conversion unavailable): {exc}"

    body = markdown.markdown(markdown_text, extensions=["tables"])
    html_text = f"<html><head><meta charset=\"utf-8\"></head><body>{body}</body></html>"
    return html_text, None


def _docx_style_exists(doc: Any, style_name: str) -> bool:
    try:
        _ = doc.styles[style_name]
        return True
    except Exception:
        return False


def _set_cell_border(cell: Any, **borders: dict[str, str]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, edge_data in borders.items():
        edge_element = tc_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_element)
        for key, value in edge_data.items():
            edge_element.set(qn(f"w:{key}"), value)


def apply_apa_header_border(table: Any) -> None:
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        _set_cell_border(
            cell,
            top={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
        )
    for cell in table.rows[-1].cells:
        _set_cell_border(
            cell,
            bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
        )


def build_docx_report(
    template_path: Path,
    output_path: Path,
    data_path: Path,
    omv_path: Path | None,
    records: Sequence[Any],
    mode: str,
    table_style: str,
    column_manifest: dict[str, Any] | None,
    timings: dict[str, float] | None,
) -> tuple[Path | None, str | None]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        return None, f"DOCX export skipped (python-docx not available): {exc}"

    if not template_path.exists():
        return None, f"DOCX export skipped (template missing): {template_path}"

    doc = Document(str(template_path))
    doc.add_paragraph("Key Summary", style="Heading 1")
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Mode: {mode}")
    doc.add_paragraph(f"Data file: {data_path}")
    doc.add_paragraph(f"OMV path: {omv_path}" if omv_path else "OMV path: not saved")

    if timings:
        doc.add_paragraph(f"Total time: {timings.get('total_seconds', 0.0):.2f}s")

    if column_manifest:
        doc.add_paragraph("Column Mapping", style="Heading 2")
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Clean Alias"
        header_cells[1].text = "Original Header"
        for alias, original in column_manifest.items():
            row_cells = table.add_row().cells
            row_cells[0].text = alias
            row_cells[1].text = str(original)
        if table_style == "apa":
            if _docx_style_exists(doc, "APATable"):
                table.style = "APATable"
            apply_apa_header_border(table)
        elif _docx_style_exists(doc, "Table Grid"):
            table.style = "Table Grid"

    successful = [record for record in records if record.status == "success"]
    table_index = 1
    for record in successful:
        doc.add_paragraph(f"{record.title} ({record.analysis_type})", style="Heading 2")
        if not record.summary_sections:
            doc.add_paragraph("Key results could not be extracted automatically.")
            continue
        for section in record.summary_sections:
            rows = section.get("rows") or []
            if not rows:
                doc.add_paragraph("No extractable rows were available.")
                continue
            if table_style == "apa":
                doc.add_paragraph(f"Table {table_index}")
                title_paragraph = doc.add_paragraph(section.get("title", ""))
                if _docx_style_exists(doc, "TableTitle"):
                    title_paragraph.style = "TableTitle"
                else:
                    for run in title_paragraph.runs:
                        run.italic = True
                table_index += 1
            else:
                doc.add_paragraph(section.get("title", ""), style="Heading 3")

            columns = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(columns))
            header_cells = table.rows[0].cells
            for idx, col in enumerate(columns):
                header_cells[idx].text = str(col)
            for row in rows:
                row_cells = table.add_row().cells
                for idx, col in enumerate(columns):
                    row_cells[idx].text = str(row.get(col, "") or "")

            if table_style == "apa":
                if _docx_style_exists(doc, "APATable"):
                    table.style = "APATable"
                apply_apa_header_border(table)
            elif _docx_style_exists(doc, "Table Grid"):
                table.style = "Table Grid"

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path, None
    except Exception as exc:
        return None, f"DOCX export failed: {exc}"


async def write_latex_bundle(html_text: str, output_path: Path) -> None:
    from jamovi.server.utils.latexify import latexify

    async def resolve_image(part: str) -> str:
        return ""

    with output_path.open("wb") as handle:
        async for _ in latexify(html_text, handle, resolve_image):
            pass


async def export_report_formats(
    markdown_text: str,
    output_base: Path,
    export_formats: Sequence[str],
) -> tuple[dict[str, str | None], list[str]]:
    warnings: list[str] = []
    paths: dict[str, str | None] = {}

    needs_html = any(fmt in export_formats for fmt in ("html", "pdf", "latex"))
    html_text = ""
    if needs_html:
        html_text, html_warning = render_markdown_html(markdown_text)
        if html_warning:
            warnings.append(html_warning)

    if "html" in export_formats:
        html_path = output_base.with_suffix(".html")
        try:
            html_path.write_text(html_text, encoding="utf-8")
            paths["html"] = str(html_path)
        except Exception as exc:
            warnings.append(f"HTML export failed: {exc}")
            paths["html"] = None

    if "latex" in export_formats:
        latex_path = output_base.with_suffix(".zip")
        try:
            await write_latex_bundle(html_text, latex_path)
            paths["latex"] = str(latex_path)
        except Exception as exc:
            warnings.append(f"LaTeX export failed: {exc}")
            paths["latex"] = None

    if "pdf" in export_formats:
        pdf_path = output_base.with_suffix(".pdf")
        try:
            import weasyprint  # type: ignore
        except Exception as exc:
            warnings.append(f"PDF export skipped (weasyprint not available): {exc}")
            paths["pdf"] = None
        else:
            try:
                weasyprint.HTML(string=html_text).write_pdf(str(pdf_path))
                paths["pdf"] = str(pdf_path)
            except Exception as exc:
                warnings.append(f"PDF export failed: {exc}")
                paths["pdf"] = None

    return paths, warnings
